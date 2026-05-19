"""
Document text extraction for all supported file formats.
Returns plain text chunks suitable for embedding in ChromaDB.
"""
import io
import mimetypes
from typing import List


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract all text from a file. Returns a single string."""
    mime, _ = mimetypes.guess_type(filename)
    name_lower = filename.lower()

    if name_lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    elif name_lower.endswith((".xlsx", ".xls")):
        return _extract_excel(file_bytes)
    elif name_lower.endswith(".csv"):
        return _extract_csv(file_bytes)
    elif name_lower.endswith((".txt", ".md")):
        return file_bytes.decode("utf-8", errors="ignore")
    elif name_lower.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
        return _extract_image_ocr(file_bytes)
    else:
        # Fallback: try to decode as text
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""


def _extract_pdf(file_bytes: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _extract_excel(file_bytes: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    rows = []
    for sheet in wb.worksheets:
        rows.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            row_str = " | ".join(str(cell) for cell in row if cell is not None)
            if row_str.strip():
                rows.append(row_str)
    return "\n".join(rows)


def _extract_csv(file_bytes: bytes) -> str:
    import csv
    text = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    return "\n".join(" | ".join(row) for row in reader)


def _extract_image_ocr(file_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(img)
    except Exception:
        return ""


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks for embedding."""
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks
